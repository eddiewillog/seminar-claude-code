"""
05_generate_report.py  (Word/.docx 버전)

요약 테이블(data/output/) 과 파일 점검 결과(data/processed/input_file_inventory.csv)
를 읽어, 재무팀 보고용 Word 문서를 생성한다.

산출:
  reports/finance_expense_report.docx

원칙:
  - 데이터에서 직접 읽은 사실만 기록한다.
  - 모르는 값은 "확인 필요" 라고 적는다.
  - 중복 인보이스는 단정하지 않고 "중복 가능성 있음" 으로 표현한다.
  - 표는 docx 의 진짜 표(add_table)로 만든다.

의존성:
  pip install python-docx
"""

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parent.parent
PROCESSED_DIR = PROJECT_ROOT / "data" / "processed"
OUTPUT_DIR = PROJECT_ROOT / "data" / "output"
REPORTS_DIR = PROJECT_ROOT / "reports"

INVENTORY_PATH = PROCESSED_DIR / "input_file_inventory.csv"
MERGED_PATH = PROCESSED_DIR / "merged_expenses.csv"
REPORT_PATH = REPORTS_DIR / "finance_expense_report.docx"


def fmt_amount(value) -> str:
    """금액을 사람이 보기 좋게 표시 (예: 1,234,567 원)."""
    if pd.isna(value):
        return "확인 필요"
    try:
        return f"{int(round(float(value))):,} 원"
    except (TypeError, ValueError):
        return "확인 필요"


def safe_read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def add_dataframe_as_table(
    doc: Document,
    df: pd.DataFrame,
    amount_cols: list[str] | None = None,
) -> None:
    """DataFrame 을 Word 표로 추가한다. 비어 있으면 안내 문구만 넣는다."""
    if df.empty:
        doc.add_paragraph("데이터 없음")
        return

    amount_cols = amount_cols or []
    cols = list(df.columns)

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"

    # 헤더
    header_cells = table.rows[0].cells
    for i, c in enumerate(cols):
        header_cells[i].text = str(c)
        for run in header_cells[i].paragraphs[0].runs:
            run.bold = True

    # 본문
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, c in enumerate(cols):
            v = row[c]
            if c in amount_cols:
                row_cells[i].text = fmt_amount(v)
            elif pd.isna(v):
                row_cells[i].text = ""
            else:
                row_cells[i].text = str(v)


def build_report() -> Document:
    today = datetime.now().strftime("%Y-%m-%d")

    inventory = safe_read_csv(INVENTORY_PATH)
    merged = safe_read_csv(MERGED_PATH)
    monthly = safe_read_csv(OUTPUT_DIR / "monthly_expense_summary.csv")
    dept = safe_read_csv(OUTPUT_DIR / "department_expense_summary.csv")
    category = safe_read_csv(OUTPUT_DIR / "category_expense_summary.csv")
    vendor_top10 = safe_read_csv(OUTPUT_DIR / "vendor_top10_summary.csv")
    duplicates = safe_read_csv(OUTPUT_DIR / "duplicate_invoice_list.csv")
    quality = safe_read_csv(OUTPUT_DIR / "data_quality_summary.csv")

    included = inventory[inventory["status"] == "include"] if not inventory.empty else inventory
    excluded = inventory[inventory["status"] == "exclude"] if not inventory.empty else inventory

    total_amount = merged["amount"].sum() if "amount" in merged.columns else 0
    total_rows = len(merged)

    doc = Document()

    # 기본 폰트 (한글 깨짐 방지용 단순 설정)
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # ---- 제목 ----
    doc.add_heading("재무 비용 분석 리포트", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"작성일: {today}").italic = True
    doc.add_paragraph("본 리포트는 합성 데이터를 기반으로 한 실습용 산출물이다.").italic = True

    # 1. 분석 개요
    doc.add_heading("1. 분석 개요", level=1)
    doc.add_paragraph(
        "프로젝트 루트의 Excel/CSV 파일을 자동으로 점검·통합한 결과를 정리한다."
    )
    doc.add_paragraph(f"분석 대상 파일 수: {len(included)}개")
    doc.add_paragraph(f"제외 파일 수: {len(excluded)}개")
    doc.add_paragraph(f"통합 거래 건수: {total_rows}건")
    doc.add_paragraph(f"통합 거래 총액: {fmt_amount(total_amount)}")

    # 2. 분석 대상 파일
    doc.add_heading("2. 분석 대상 파일 목록", level=1)
    if included.empty:
        doc.add_paragraph("분석 대상 파일이 없다.")
    else:
        view = included[["file_name", "file_type", "row_count", "note"]].copy()
        add_dataframe_as_table(doc, view)

    # 3. 제외 파일
    doc.add_heading("3. 제외 파일 목록", level=1)
    if excluded.empty:
        doc.add_paragraph("제외된 파일이 없다.")
    else:
        view = excluded[["file_name", "file_type", "note"]].copy()
        add_dataframe_as_table(doc, view)

    # 4. 전체 비용 합계
    doc.add_heading("4. 전체 비용 합계", level=1)
    doc.add_paragraph(f"총 거래 건수: {total_rows}건")
    doc.add_paragraph(f"총 결제 금액: {fmt_amount(total_amount)}")

    # 5. 월별 비용 요약
    doc.add_heading("5. 월별 비용 요약", level=1)
    add_dataframe_as_table(doc, monthly, amount_cols=["total_amount"])

    # 6. 부서별 비용 요약
    doc.add_heading("6. 부서별 비용 요약", level=1)
    add_dataframe_as_table(doc, dept, amount_cols=["total_amount"])

    # 7. 비용 항목별 요약
    doc.add_heading("7. 비용 항목별 요약", level=1)
    add_dataframe_as_table(doc, category, amount_cols=["total_amount"])

    # 8. 거래처 Top 10
    doc.add_heading("8. 거래처 Top 10", level=1)
    add_dataframe_as_table(doc, vendor_top10, amount_cols=["total_amount"])

    # 9. 데이터 품질 점검 결과
    doc.add_heading("9. 데이터 품질 점검 결과", level=1)
    add_dataframe_as_table(doc, quality)

    # 10. 중복 인보이스
    doc.add_heading("10. 중복 invoice_no 목록 (중복 가능성 있음)", level=1)
    if duplicates.empty:
        doc.add_paragraph("중복으로 추정되는 인보이스가 없다.")
    else:
        add_dataframe_as_table(doc, duplicates, amount_cols=["amount"])

    # 11. 확인 필요 사항
    doc.add_heading("11. 재무팀 확인 필요 사항", level=1)
    if not quality.empty:
        issues = quality[quality["count"].astype(int) > 0]
        issues = issues[issues["issue"] != "총 거래 건수"]
        if issues.empty:
            doc.add_paragraph("별도의 데이터 품질 이슈가 발견되지 않았다.")
        else:
            for _, r in issues.iterrows():
                doc.add_paragraph(
                    f"{r['issue']}: {int(r['count'])}건 → 원본 파일 재확인 필요",
                    style="List Bullet",
                )
    else:
        doc.add_paragraph("품질 요약이 비어 있다. 확인 필요.")

    # 12. 다음 액션
    doc.add_heading("12. 다음 액션", level=1)
    for line in [
        "제외 파일 목록을 재무팀 담당자와 함께 검토한다.",
        "날짜/금액 변환 실패 행은 원본 파일에서 표기를 확인한다.",
        "중복 invoice_no 목록은 동일 거래 중복 입력 여부를 확인한다.",
        "확정된 결과는 별도 보고용 양식에 정리해 공유한다.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    return doc


def main() -> None:
    print("[05] Word 리포트 생성을 시작한다.")
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(REPORT_PATH)
    print(f"  저장됨: {REPORT_PATH.relative_to(PROJECT_ROOT)}")
    print("[05] 완료. 다음 단계는 06_upload_to_google_docs.py 입니다.")


if __name__ == "__main__":
    main()

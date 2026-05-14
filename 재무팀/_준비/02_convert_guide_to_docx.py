"""
02_convert_guide_to_docx.py  (워크숍 자료 빌드용)

재무팀/실습_가이드.md 를 워드 파일(.docx)로 변환한다.
트레이니가 가이드를 PDF나 워드 형태로 옆에 띄워두고 보는 용도다.

산출:
  ../실습_가이드.docx

지원하는 마크다운 요소:
  - 헤딩 (#, ##, ###)
  - 단락
  - 코드 블록 ```
  - 표 | a | b |
  - 인용 블록 > (내부의 코드 블록 포함)
  - 글머리표 (-), 번호 매기기 (1.)
  - 인라인 굵게(**), 인라인 코드(`)
  - 수평선 (---)

가이드 문서를 수정한 뒤 이 스크립트를 다시 실행하면 .docx 도 같이 갱신된다.

의존성:
  pip install python-docx
"""

from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
GUIDE_MD = PROJECT_ROOT / "실습_가이드.md"
GUIDE_DOCX = PROJECT_ROOT / "실습_가이드.docx"


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def setup_styles(doc: Document) -> None:
    """한글 폰트를 기본으로 적용."""
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10)


def add_inline_paragraph(
    doc: Document,
    text: str,
    style: str | None = None,
    italic: bool = False,
    left_indent: float | None = None,
):
    """인라인 마크다운(**bold**, `code`) 을 살린 단락을 추가한다."""
    p = doc.add_paragraph()
    if style:
        p.style = style
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)

    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run.text = part
        if italic:
            run.italic = True
    return p


def add_code_block(doc: Document, code_text: str, left_indent: float = 0.3) -> None:
    """코드 블록을 모노스페이스 단락으로 추가한다."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(left_indent)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _strip_inline_marks(text: str) -> str:
    """표 셀 안의 **bold**, `code` 표기를 단순 텍스트로 정리."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = text.replace("`", "")
    return text


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    """`|` 로 구성된 마크다운 표를 docx 표로 변환한다."""
    if len(table_lines) < 2:
        return

    def split_row(line: str) -> list[str]:
        # 양쪽 끝의 | 를 떼고 분할
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        return [c.strip() for c in line.split("|")]

    header = split_row(table_lines[0])
    # table_lines[1] 은 구분선 (| --- | --- |) → 스킵
    body_rows = [split_row(ln) for ln in table_lines[2:]]

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"

    # 헤더
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(_strip_inline_marks(h))
        run.bold = True

    # 본문
    for row in body_rows:
        row_cells = table.add_row().cells
        for i in range(len(header)):
            val = row[i] if i < len(row) else ""
            row_cells[i].text = _strip_inline_marks(val)


def render_blockquote(doc: Document, quote_lines: list[str]) -> None:
    """
    > 로 시작하는 인용 블록을 처리한다.
    내부에 ``` 코드 블록이 있으면 그건 코드 단락으로 따로 그린다.
    """
    # 앞의 "> " 또는 ">" 를 제거해 순수 내용만 남긴다.
    stripped: list[str] = []
    for ln in quote_lines:
        if ln.startswith("> "):
            stripped.append(ln[2:])
        elif ln.startswith(">"):
            stripped.append(ln[1:].lstrip())
        else:
            stripped.append(ln)

    j = 0
    n = len(stripped)
    while j < n:
        cur = stripped[j]

        if cur.startswith("```"):
            # 인용 안의 코드 블록 — 닫는 ``` 까지 모은다.
            end = j + 1
            while end < n and not stripped[end].startswith("```"):
                end += 1
            code_text = "\n".join(stripped[j + 1 : end])
            add_code_block(doc, code_text, left_indent=0.5)
            j = end + 1
        elif cur.strip() == "":
            j += 1
        else:
            # 본문 텍스트 — 다음 빈 줄 또는 코드 블록 시작 전까지 합친다.
            text_lines = [cur]
            j += 1
            while j < n and stripped[j].strip() != "" and not stripped[j].startswith("```"):
                text_lines.append(stripped[j])
                j += 1
            add_inline_paragraph(
                doc,
                " ".join(text_lines),
                italic=True,
                left_indent=0.25,
            )


def convert(md_text: str) -> Document:
    doc = Document()
    setup_styles(doc)

    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # 헤딩
        m = re.match(r"^(#{1,3}) (.+)$", line)
        if m:
            level = len(m.group(1)) - 1
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # 수평선
        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # 코드 블록 (인용 밖)
        if line.startswith("```"):
            end = i + 1
            while end < n and not lines[end].startswith("```"):
                end += 1
            code_text = "\n".join(lines[i + 1 : end])
            add_code_block(doc, code_text)
            i = end + 1
            continue

        # 표
        if line.startswith("|"):
            table_lines = []
            while i < n and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        # 인용 블록
        if line.startswith(">"):
            quote_lines = []
            while i < n:
                if lines[i].startswith(">"):
                    quote_lines.append(lines[i])
                    i += 1
                elif lines[i].strip() == "" and i + 1 < n and lines[i + 1].startswith(">"):
                    # 인용 블록 안의 빈 줄
                    quote_lines.append(lines[i])
                    i += 1
                else:
                    break
            render_blockquote(doc, quote_lines)
            continue

        # 글머리표
        if line.startswith("- "):
            add_inline_paragraph(doc, line[2:], style="List Bullet")
            i += 1
            continue

        # 번호 매기기
        m = re.match(r"^\d+\. (.+)$", line)
        if m:
            add_inline_paragraph(doc, m.group(1), style="List Number")
            i += 1
            continue

        # 빈 줄
        if line.strip() == "":
            i += 1
            continue

        # 일반 단락
        add_inline_paragraph(doc, line)
        i += 1

    return doc


def main() -> None:
    print("[02] 가이드 마크다운을 워드로 변환한다.")
    if not GUIDE_MD.exists():
        raise FileNotFoundError(f"가이드 파일이 없다: {GUIDE_MD}")

    md_text = GUIDE_MD.read_text(encoding="utf-8")
    doc = convert(md_text)
    doc.save(GUIDE_DOCX)
    print(f"  저장됨: {GUIDE_DOCX.relative_to(PROJECT_ROOT.parent)}")
    print("[02] 완료.")


if __name__ == "__main__":
    main()
